"""
4.4 可视化与报告导出
- APA 表格
- lavaan 语法导出
"""
from io import BytesIO
from typing import Any
from datetime import datetime
import re
import hashlib

import pandas as pd
from docx import Document
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

from app.api.errors import ErrorCode, api_error
from app.services.data_parser import get_data

router = APIRouter()
DOCX_MAX_TABLE_ROWS = 200
EXPORT_TEMPLATE_REGISTRY: dict[str, dict[str, str]] = {
    "journal_minimal_cn": {
        "name": "期刊模板（最小版）",
        "description": "最小可用期刊导出模板，保持现有字段契约并补充模板追溯信息。",
    }
}


class ExportPayload(BaseModel):
    """导出请求载荷"""
    data_key: str
    export_template: str | None = None
    lavaan_syntax: str
    fit_indices: dict[str, Any] | None = None
    estimator: str | None = None
    missing_strategy: str | None = None
    fit_ran_at: str | None = None
    mi_adoption_count: int | None = None
    mi_ignored_count: int | None = None
    mi_risk_filter: str | None = None
    # E3-6：导出可追溯/一致性（前端提供的“版本/采纳摘要”）
    syntax_tag: str | None = None
    last_fit_syntax_tag: str | None = None
    fit_stale: bool | None = None
    mi_adoption_bucket_tag: str | None = None
    mi_adoption_first_at: str | None = None
    mi_adoption_last_at: str | None = None
    mi_adoption_last_id: str | None = None
    # E3-6：可选门禁（默认拒绝导出 stale 结果；允许时需显式声明）
    allow_stale_export: bool | None = None
    estimates: list[dict[str, Any]] | None = None
    bootstrap: dict[str, Any] | None = None
    moderation: dict[str, Any] | None = None
    moderated_mediation: dict[str, Any] | None = None
    invariance: dict[str, Any] | None = None
    invariance_series: dict[str, Any] | None = None
    # Wave 2 / P1：不变性闭环补强（前端提供的“lite/strict 摘要与可复制段落”）
    invariance_series_summary_lines: list[str] | None = None
    invariance_series_conclusion: dict[str, Any] | None = None
    invariance_series_report: dict[str, Any] | None = None
    # Wave 2 / P1：路径系数显著性汇总 + 论文式结果摘要（前端生成，导出可追溯）
    path_summary_rows: list[dict[str, Any]] | None = None
    path_summary_lines: list[str] | None = None
    path_summary_report_text: str | None = None
    # 1.3：综合结果摘要（拟合 + 路径 + 可选 Bootstrap/调节，前端拼装）
    integrated_paper_paragraph: str | None = None
    model_comparison: dict[str, Any] | None = None
    mga_structural: dict[str, Any] | None = None
    # 1.3：潜变量交互 MVP 探测结果（可选）
    latent_interaction: dict[str, Any] | None = None
    project_name: str | None = None


def _safe_text(v: Any) -> str:
    if v is None:
        return "-"
    return str(v)

def _safe_na(v: Any) -> str:
    """
    导出用：空值统一写 NA（比 "-" 更适合论文表格/Excel 后处理）。
    """
    if v is None:
        return "NA"
    if v == "-":
        return "NA"
    return str(v)


def _truncate_docx_rows(rows: list[Any], *, limit: int = DOCX_MAX_TABLE_ROWS) -> tuple[list[Any], int]:
    if limit <= 0:
        return [], len(rows or [])
    safe_rows = list(rows or [])
    if len(safe_rows) <= limit:
        return safe_rows, 0
    return safe_rows[:limit], len(safe_rows) - limit


def _raise_export_error(
    status_code: int, message: str, code: str, *, hint: str | None = None
) -> None:
    raise api_error(status_code, code=code, message=message, hint=hint)


def _sanitize_project_name(name: str | None) -> str:
    raw = (name or "OpenSEM").strip()
    cleaned = re.sub(r"[^A-Za-z0-9_\-]+", "_", raw)
    return cleaned[:48] or "OpenSEM"

def _normalize_syntax_for_fingerprint(syntax: str) -> str:
    lines: list[str] = []
    for raw in (syntax or "").splitlines():
        s = raw.strip()
        if not s:
            continue
        lines.append(s)
    return "\n".join(lines)


def _syntax_fingerprint(syntax: str) -> str:
    norm = _normalize_syntax_for_fingerprint(syntax)
    h = hashlib.sha256(norm.encode("utf-8")).hexdigest()
    return h[:12]


def _build_invariance_series_summary_df(
    *,
    summary_lines: list[str] | None,
    conclusion: dict[str, Any] | None,
) -> pd.DataFrame | None:
    """
    导出用：把前端生成的“不变性摘要/结论”落为可读表格。
    目标：Excel/Word 可追溯，可复制（而不是只留在 UI）。
    """
    lines = [str(x) for x in (summary_lines or []) if str(x).strip()]
    if not lines and not (isinstance(conclusion, dict) and conclusion):
        return None

    rows: list[dict[str, Any]] = []
    if isinstance(conclusion, dict) and conclusion:
        strict = conclusion.get("strict") if isinstance(conclusion.get("strict"), dict) else {}
        lite = conclusion.get("lite") if isinstance(conclusion.get("lite"), dict) else {}
        rows.append(
            {
                "type": "conclusion",
                "scope": "lite",
                "level": _safe_text(lite.get("level")),
                "ok": _safe_text(lite.get("ok")),
                "note": _safe_text(lite.get("note")),
            }
        )
        rows.append(
            {
                "type": "conclusion",
                "scope": "strict",
                "level": _safe_text(strict.get("level")),
                "ok": _safe_text(strict.get("ok")),
                "note": _safe_text(strict.get("note")),
            }
        )

    for i, line in enumerate(lines, start=1):
        rows.append({"type": "summary_line", "scope": "all", "level": "", "ok": "", "note": line, "order": i})

    df = pd.DataFrame(rows)
    # 固定列顺序
    cols = ["type", "scope", "level", "ok", "note", "order"]
    return df.reindex(columns=cols)


def _build_invariance_series_report_df(report: dict[str, Any] | None) -> pd.DataFrame | None:
    """
    导出用：把前端生成的“不变性报告模板（lite/strict）”落为可复制的逐行表格。
    目标：Excel 可追溯，可复制（与 docx 保持一致信息）。
    """
    if not isinstance(report, dict) or not report:
        return None

    lite_text = str(report.get("lite") or "").strip()
    strict_text = str(report.get("strict") or "").strip()
    if not lite_text and not strict_text:
        return None

    rows: list[dict[str, Any]] = []
    if lite_text:
        for i, line in enumerate(lite_text.splitlines(), start=1):
            rows.append({"type": "report_line", "scope": "lite", "text": line, "order": i})
    if strict_text:
        for i, line in enumerate(strict_text.splitlines(), start=1):
            rows.append({"type": "report_line", "scope": "strict", "text": line, "order": i})

    df = pd.DataFrame(rows)
    cols = ["type", "scope", "text", "order"]
    return df.reindex(columns=cols)


def _build_path_summary_df(
    *,
    rows: list[dict[str, Any]] | None,
    summary_lines: list[str] | None,
    report_text: str | None,
) -> pd.DataFrame | None:
    """
    导出用：把前端生成的“路径显著性汇总表 + 摘要段落/要点”落为可读表格。
    目标：Excel/Word 可追溯，可复制（而不是只留在 UI）。
    """
    path_rows = [r for r in (rows or []) if isinstance(r, dict)]
    lines = [str(x) for x in (summary_lines or []) if str(x).strip()]
    report = str(report_text or "").strip()
    if not path_rows and not lines and not report:
        return None

    out_rows: list[dict[str, Any]] = []
    for r in path_rows:
        out_rows.append(
            {
                "type": "path",
                "predictor": _safe_text(r.get("predictor")),
                "outcome": _safe_text(r.get("outcome")),
                "op": _safe_text(r.get("op") or "~"),
                "beta": _safe_na(r.get("beta")),
                "p_value": _safe_na(r.get("p_value")),
                "significant": _safe_text(r.get("significant")),
                "direction": _safe_text(r.get("direction")),
                "note": _safe_text(r.get("note")),
                "text": "",
                "order": "",
            }
        )

    if lines:
        for i, line in enumerate(lines, start=1):
            out_rows.append(
                {
                    "type": "summary_line",
                    "predictor": "",
                    "outcome": "",
                    "op": "",
                    "beta": "",
                    "p_value": "",
                    "significant": "",
                    "direction": "",
                    "note": "",
                    "text": line,
                    "order": i,
                }
            )

    if report:
        for i, line in enumerate(report.splitlines(), start=1):
            out_rows.append(
                {
                    "type": "report_line",
                    "predictor": "",
                    "outcome": "",
                    "op": "",
                    "beta": "",
                    "p_value": "",
                    "significant": "",
                    "direction": "",
                    "note": "",
                    "text": line,
                    "order": i,
                }
            )

    df = pd.DataFrame(out_rows)
    cols = ["type", "predictor", "outcome", "op", "beta", "p_value", "significant", "direction", "note", "text", "order"]
    return df.reindex(columns=cols)


def _build_mga_structural_dfs(mga_structural: dict[str, Any] | None):
    """
    导出用：兼容 Wave 2（单路径）与 Wave 3（层级×多路径逐步比较）。

    - Wave 2（旧）：
      {
        group_estimates: [...],
        comparison: {...},
        path: {predictor,outcome},
      }

    - Wave 3（新）：
      {
        items: [
          { level, path, group_estimates, comparison, ... },
          ...
        ]
      }

    返回：(mga_path_df, mga_lrt_df)，任意一项可为 None。
    """
    if not isinstance(mga_structural, dict):
        return None, None

    # 新结构：items（level×path）
    items = mga_structural.get("items")
    if isinstance(items, list) and items:
        rows_path: list[dict[str, Any]] = []
        rows_lrt: list[dict[str, Any]] = []
        for it in items:
            if not isinstance(it, dict):
                continue
            level = _safe_text(it.get("level"))
            p = it.get("path") if isinstance(it.get("path"), dict) else {}
            predictor = _safe_text(p.get("predictor"))
            outcome = _safe_text(p.get("outcome"))

            for ge in it.get("group_estimates") or []:
                if not isinstance(ge, dict):
                    continue
                coef = ge.get("coef") if isinstance(ge.get("coef"), dict) else {}
                rows_path.append(
                    {
                        "level": level,
                        "predictor": predictor,
                        "outcome": outcome,
                        "group": _safe_text(ge.get("group")),
                        "estimate": _safe_na(ge.get("estimate") if ge.get("estimate") is not None else coef.get("estimate")),
                        "std_all": _safe_na(ge.get("std_all") if ge.get("std_all") is not None else coef.get("std_all")),
                        "p_value": _safe_na(ge.get("p_value") if ge.get("p_value") is not None else coef.get("p_value")),
                        "note": _safe_text(ge.get("error") if ge.get("success") is False else (ge.get("note") or "")),
                    }
                )

            comp = it.get("comparison") if isinstance(it.get("comparison"), dict) else None
            if comp:
                rows_lrt.append(
                    {
                        "level": level,
                        "predictor": predictor,
                        "outcome": outcome,
                        "from": comp.get("from"),
                        "to": comp.get("to"),
                        "ok": comp.get("ok"),
                        "chi2_diff": _safe_na(comp.get("chi2_diff")),
                        "df_diff": _safe_na(comp.get("df_diff")),
                        "p_value": _safe_na(comp.get("p_value")),
                        "note": comp.get("note"),
                    }
                )

        mga_path_df = pd.DataFrame(rows_path) if rows_path else None
        if mga_path_df is not None:
            cols = ["level", "predictor", "outcome", "group", "estimate", "std_all", "p_value", "note"]
            mga_path_df = mga_path_df.reindex(columns=cols)

        mga_lrt_df = pd.DataFrame(rows_lrt) if rows_lrt else None
        if mga_lrt_df is not None:
            cols = ["level", "predictor", "outcome", "from", "to", "ok", "chi2_diff", "df_diff", "p_value", "note"]
            mga_lrt_df = mga_lrt_df.reindex(columns=cols)

        return mga_path_df, mga_lrt_df

    # 旧结构：单路径
    mga_path_df = None
    mga_lrt_df = None

    if mga_structural.get("group_estimates"):
        try:
            rows = []
            for it in mga_structural.get("group_estimates") or []:
                if not isinstance(it, dict):
                    continue
                coef = it.get("coef") if isinstance(it.get("coef"), dict) else {}
                rows.append(
                    {
                        "group": _safe_text(it.get("group")),
                        "estimate": _safe_na(it.get("estimate") if it.get("estimate") is not None else coef.get("estimate")),
                        "std_all": _safe_na(it.get("std_all") if it.get("std_all") is not None else coef.get("std_all")),
                        "p_value": _safe_na(it.get("p_value") if it.get("p_value") is not None else coef.get("p_value")),
                        "note": _safe_text(it.get("error") if it.get("success") is False else (it.get("note") or "")),
                    }
                )
            mga_path_df = pd.DataFrame(rows)
            cols = ["group", "estimate", "std_all", "p_value", "note"]
            mga_path_df = mga_path_df.reindex(columns=cols)
        except Exception:
            mga_path_df = pd.DataFrame([{"error": "mga_structural group_estimates parse failed"}])

    if mga_structural.get("comparison"):
        try:
            c = mga_structural.get("comparison") or {}
            mga_lrt_df = pd.DataFrame(
                [
                    {
                        "from": c.get("from"),
                        "to": c.get("to"),
                        "ok": c.get("ok"),
                        "chi2_diff": _safe_na(c.get("chi2_diff")),
                        "df_diff": _safe_na(c.get("df_diff")),
                        "p_value": _safe_na(c.get("p_value")),
                        "note": c.get("note"),
                    }
                ]
            )
            cols = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "note"]
            mga_lrt_df = mga_lrt_df.reindex(columns=cols)
        except Exception:
            mga_lrt_df = pd.DataFrame([{"error": "mga_structural comparison parse failed"}])

    return mga_path_df, mga_lrt_df


def _flatten_bootstrap_items(items: list[dict[str, Any]], *, covariates_text: str | None = None) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    cov_text = _safe_text(covariates_text)
    for raw in items or []:
        row = raw or {}
        ci = row.get("ci") or {}
        percentile = ci.get("percentile") or {}
        bc = ci.get("bc") or {}
        sequence = row.get("sequence")
        mediators = row.get("mediators")
        component_paths = row.get("component_paths")
        rows.append(
            {
                "effect_type": _safe_text(row.get("effect_type")),
                "label": _safe_text(row.get("label")),
                "path": _safe_text(row.get("path_label") or (" → ".join(sequence) if isinstance(sequence, list) else None)),
                "x": _safe_text(row.get("x")),
                "y": _safe_text(row.get("y")),
                "mediators": _safe_text(" → ".join(mediators) if isinstance(mediators, list) and mediators else None),
                "components": _safe_text(" ; ".join(component_paths) if isinstance(component_paths, list) and component_paths else None),
                "covariates": cov_text,
                "indirect_point": _safe_na(row.get("indirect_point")),
                "ci_level": _safe_na(ci.get("ci_level")),
                "percentile_lo": _safe_na(percentile.get("lo")),
                "percentile_hi": _safe_na(percentile.get("hi")),
                "bc_lo": _safe_na(bc.get("lo")),
                "bc_hi": _safe_na(bc.get("hi")),
                "n_boot_valid": _safe_na(ci.get("n_boot_valid")),
                "note": _safe_text(row.get("note")),
            }
        )
    return pd.DataFrame(rows)


def _flatten_latent_interaction_rows(payload: dict[str, Any]) -> pd.DataFrame:
    """潜变量交互 MVP 探测结果：matching / structural 分区导出。"""
    rows: list[dict[str, Any]] = []
    if not isinstance(payload, dict):
        return pd.DataFrame(rows)
    b = payload.get("boundary") or {}
    base = {
        "y": _safe_text(payload.get("y")),
        "f1": _safe_text(payload.get("f1")),
        "f2": _safe_text(payload.get("f2")),
        "n_used": _safe_na(payload.get("n_used")),
        "boundary_semopy": _safe_text(b.get("semopy")),
        "boundary_lavaan": _safe_text(b.get("lavaan")),
    }
    for section, key in (("matching", "matching"), ("structural_for_y", "structural_for_y")):
        for raw in payload.get(key) or []:
            row = raw or {}
            rows.append(
                {
                    **base,
                    "section": section,
                    "lval": _safe_text(row.get("lval")),
                    "op": _safe_text(row.get("op")),
                    "rval": _safe_text(row.get("rval")),
                    "estimate": _safe_na(row.get("estimate")),
                    "est_std": _safe_na(row.get("est_std")),
                    "p_value": _safe_na(row.get("p_value")),
                }
            )
    return pd.DataFrame(rows)


def _flatten_moderation_rows(payload: dict[str, Any]) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    if not isinstance(payload, dict):
        return pd.DataFrame(rows)

    moderator = payload.get("moderator") or {}
    covs = payload.get("covariates")
    cov_text = _safe_text(", ".join(covs) if isinstance(covs, list) and covs else None)
    for row in payload.get("coefficients") or []:
        rows.append(
            {
                "section": "coefficients",
                "term": _safe_text((row or {}).get("term")),
                "estimate": _safe_na((row or {}).get("estimate")),
                "std_err": _safe_na((row or {}).get("std_err")),
                "z_value": _safe_na((row or {}).get("z_value")),
                "p_value": _safe_na((row or {}).get("p_value")),
                "group": "NA",
                "slope_x": "NA",
                "note": "NA",
                "moderator_type": _safe_text(moderator.get("type")),
                "reference_group": _safe_text(moderator.get("reference_group")),
                "covariates": cov_text,
            }
        )
    for row in payload.get("simple_slopes") or []:
        rows.append(
            {
                "section": "simple_slopes",
                "term": _safe_text(payload.get("predictor")),
                "estimate": "NA",
                "std_err": "NA",
                "z_value": "NA",
                "p_value": "NA",
                "group": _safe_text((row or {}).get("group")),
                "slope_x": _safe_na((row or {}).get("slope_x")),
                "note": _safe_text((row or {}).get("note")),
                "moderator_type": _safe_text(moderator.get("type")),
                "reference_group": _safe_text(moderator.get("reference_group")),
                "covariates": cov_text,
            }
        )
    return pd.DataFrame(rows)


def _flatten_moderated_mediation_rows(payload: dict[str, Any]) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    if not isinstance(payload, dict):
        return pd.DataFrame(rows)

    covs = payload.get("covariates")
    cov_text = _safe_text(", ".join(covs) if isinstance(covs, list) and covs else None)
    base = {
        "x": _safe_text(payload.get("x")),
        "m": _safe_text(payload.get("m")),
        "y": _safe_text(payload.get("y")),
        "w": _safe_text(payload.get("w")),
        "w_type": _safe_text(payload.get("w_type")),
        "covariates": cov_text,
        "model": _safe_text(payload.get("model")),
    }
    for item in payload.get("conditional_indirect") or []:
        ci = (item or {}).get("ci") or {}
        perc = (ci or {}).get("percentile") or {}
        bc = (ci or {}).get("bc") or {}
        rows.append(
            {
                **base,
                "section": "conditional_indirect",
                "label": _safe_text((item or {}).get("label")),
                "w_value": _safe_na((item or {}).get("w_value")),
                "group_note": _safe_text((item or {}).get("note")),
                "effect_point": _safe_na((item or {}).get("indirect_point")),
                "ci_level": _safe_na(ci.get("ci_level")),
                "percentile_lo": _safe_na(perc.get("lo")),
                "percentile_hi": _safe_na(perc.get("hi")),
                "bc_lo": _safe_na(bc.get("lo")),
                "bc_hi": _safe_na(bc.get("hi")),
                "n_boot_valid": _safe_na(ci.get("n_boot_valid")),
            }
        )
    idx = payload.get("index_moderated_mediation") or {}
    if isinstance(idx, dict) and idx:
        ci = idx.get("ci") or {}
        perc = (ci or {}).get("percentile") or {}
        bc = (ci or {}).get("bc") or {}
        rows.append(
            {
                **base,
                "section": "index_moderated_mediation",
                "label": _safe_text(idx.get("contrast")),
                "w_value": "NA",
                "effect_point": _safe_na(idx.get("point")),
                "ci_level": _safe_na(ci.get("ci_level")),
                "percentile_lo": _safe_na(perc.get("lo")),
                "percentile_hi": _safe_na(perc.get("hi")),
                "bc_lo": _safe_na(bc.get("lo")),
                "bc_hi": _safe_na(bc.get("hi")),
                "n_boot_valid": _safe_na(ci.get("n_boot_valid")),
            }
        )
    return pd.DataFrame(rows)


def _flatten_invariance_estimates(items: list[dict[str, Any]]) -> pd.DataFrame:
    """
    导出用：把 invariance_configural 的各组参数表（estimates）展开为单表，便于 Excel 后处理与论文撰写。
    """
    rows: list[dict[str, Any]] = []
    for it in items or []:
        if not isinstance(it, dict):
            continue
        group = _safe_text(it.get("group"))
        est = it.get("estimates") or []
        if not isinstance(est, list) or not est:
            continue
        for r in est:
            if not isinstance(r, dict):
                continue
            rows.append(
                {
                    "group": group,
                    "lval": _safe_text(r.get("lval")),
                    "op": _safe_text(r.get("op")),
                    "rval": _safe_text(r.get("rval")),
                    "estimate": _safe_na(r.get("estimate")),
                    "est_std": _safe_na(r.get("est_std")),
                    "std_err": _safe_na(r.get("std_err")),
                    "z_value": _safe_na(r.get("z_value")),
                    "p_value": _safe_na(r.get("p_value")),
                }
            )
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    cols = ["group", "lval", "op", "rval", "estimate", "est_std", "std_err", "z_value", "p_value"]
    return df.reindex(columns=cols)


def _build_filename(payload: ExportPayload, suffix: str) -> str:
    project = _sanitize_project_name(payload.project_name)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{project}_{stamp}_{suffix}"


def _validate_export_payload(payload: ExportPayload):
    df = get_data(payload.data_key)
    if df is None:
        _raise_export_error(400, "data_key 无效或已过期，请先重新上传数据", "EXPORT_DATA_KEY_INVALID")

    syntax = payload.lavaan_syntax.strip()
    if not syntax:
        _raise_export_error(400, "lavaan_syntax 不能为空，请先在建模页生成语法", "EXPORT_SYNTAX_EMPTY")

    # 可选门禁：当 fit_stale=true 时默认拒绝导出（除非 allow_stale_export=true）
    if payload.fit_stale is True and payload.allow_stale_export is not True:
        _raise_export_error(
            400,
            "检测到语法已在上次估算后发生变化（fit_stale=true）。为避免导出旧结果误导，后端已拒绝导出。请先重新估算后再导出；或在导出请求中显式设置 allow_stale_export=true 以强制导出。",
            "EXPORT_FIT_STALE_BLOCKED",
        )

    return df, syntax


def _resolve_export_template(payload: ExportPayload) -> tuple[str, dict[str, str]]:
    template_key = (payload.export_template or "journal_minimal_cn").strip()
    template = EXPORT_TEMPLATE_REGISTRY.get(template_key)
    if template is None:
        _raise_export_error(
            400,
            f"不支持的导出模板：{template_key}。当前仅支持 journal_minimal_cn。",
            "EXPORT_TEMPLATE_UNSUPPORTED",
        )
    return template_key, template


@router.post("/apa-table")
def export_apa_table(payload: ExportPayload):
    """
    导出 APA 7th 三线表 (.xlsx)
    """
    df, syntax = _validate_export_payload(payload)
    template_key, template = _resolve_export_template(payload)

    fit = payload.fit_indices or {}
    fit_rows = [
        {"Indicator": "chi2", "Value": _safe_text(fit.get("chi2"))},
        {"Indicator": "chi2_df", "Value": _safe_text(fit.get("chi2_df"))},
        {"Indicator": "RMSEA", "Value": _safe_text(fit.get("rmsea"))},
        {"Indicator": "SRMR", "Value": _safe_text(fit.get("srmr"))},
        {"Indicator": "CFI", "Value": _safe_text(fit.get("cfi"))},
        {"Indicator": "TLI", "Value": _safe_text(fit.get("tli"))},
        {"Indicator": "status", "Value": _safe_text(fit.get("status"))},
    ]
    fit_df = pd.DataFrame(fit_rows)

    syntax_lines = [line.strip() for line in syntax.splitlines() if line.strip()]
    syntax_df = pd.DataFrame({"Lavaan Syntax": syntax_lines})

    est_rows = payload.estimates or []
    est_df = pd.DataFrame(est_rows)
    if not est_df.empty:
        rename_map = {
            "lval": "LHS",
            "op": "Op",
            "rval": "RHS",
            "estimate": "Estimate",
            "est_std": "Beta(Std)",
            "std_err": "S.E.",
            "z_value": "C.R.",
            "p_value": "p",
        }
        est_df = est_df.rename(columns={k: v for k, v in rename_map.items() if k in est_df.columns})
    meta_df = pd.DataFrame(
        [
            {"Item": "N Rows (raw)", "Value": int(len(df))},
            {"Item": "N Cols", "Value": int(len(df.columns))},
            {"Item": "Export Template Key", "Value": template_key},
            {"Item": "Export Template Name", "Value": template.get("name")},
            {"Item": "Export Template Description", "Value": template.get("description")},
            {"Item": "Estimator", "Value": _safe_text(payload.estimator)},
            {"Item": "Missing Strategy", "Value": _safe_text(payload.missing_strategy)},
            {"Item": "Fit Ran At (client)", "Value": _safe_text(payload.fit_ran_at)},
            {"Item": "Syntax Tag (client)", "Value": _safe_text(payload.syntax_tag)},
            {"Item": "Last Fit Syntax Tag (client)", "Value": _safe_text(payload.last_fit_syntax_tag)},
            {"Item": "Fit Stale (client)", "Value": _safe_text(payload.fit_stale)},
            {"Item": "MI Adoption Count", "Value": _safe_text(payload.mi_adoption_count)},
            {"Item": "MI Ignored Count", "Value": _safe_text(payload.mi_ignored_count)},
            {"Item": "MI Risk Filter", "Value": _safe_text(payload.mi_risk_filter)},
            {"Item": "MI Adoption Bucket Tag", "Value": _safe_text(payload.mi_adoption_bucket_tag)},
            {"Item": "MI Adoption First At", "Value": _safe_text(payload.mi_adoption_first_at)},
            {"Item": "MI Adoption Last At", "Value": _safe_text(payload.mi_adoption_last_at)},
            {"Item": "MI Adoption Last ID", "Value": _safe_text(payload.mi_adoption_last_id)},
            {"Item": "Syntax Fingerprint (sha256/12)", "Value": _syntax_fingerprint(syntax)},
            {"Item": "Note", "Value": "OpenSEM export (Phase 1/3). Optional: Bootstrap, Invariance, Invariance_Series_*"},
        ]
    )

    bootstrap_df = None
    if isinstance(payload.bootstrap, dict) and payload.bootstrap.get("items"):
        try:
            covs = payload.bootstrap.get("covariates")
            cov_text = ", ".join([str(x).strip() for x in covs]) if isinstance(covs, list) and covs else None
            bootstrap_df = _flatten_bootstrap_items(payload.bootstrap.get("items") or [], covariates_text=cov_text)
        except Exception:
            bootstrap_df = pd.DataFrame([{"error": "bootstrap payload parse failed"}])

    moderation_df = None
    if isinstance(payload.moderation, dict) and (
        payload.moderation.get("coefficients") or payload.moderation.get("simple_slopes")
    ):
        try:
            moderation_df = _flatten_moderation_rows(payload.moderation)
        except Exception:
            moderation_df = pd.DataFrame([{"error": "moderation payload parse failed"}])

    moderated_mediation_df = None
    if isinstance(payload.moderated_mediation, dict) and (
        payload.moderated_mediation.get("conditional_indirect") or payload.moderated_mediation.get("index_moderated_mediation")
    ):
        try:
            moderated_mediation_df = _flatten_moderated_mediation_rows(payload.moderated_mediation)
        except Exception:
            moderated_mediation_df = pd.DataFrame([{"error": "moderated_mediation payload parse failed"}])

    invariance_df = None
    invariance_estimates_df = None
    if isinstance(payload.invariance, dict) and payload.invariance.get("items"):
        try:
            items = payload.invariance.get("items") or []
            # flatten a bit for readability
            rows = []
            for it in items:
                fit_i = (it or {}).get("fit_indices") or {}
                rows.append(
                    {
                        "group": (it or {}).get("group"),
                        "n_used": (it or {}).get("n_used"),
                        "status": fit_i.get("status"),
                        "chi2_df": fit_i.get("chi2_df"),
                        "rmsea": fit_i.get("rmsea"),
                        "cfi": fit_i.get("cfi"),
                        "tli": fit_i.get("tli"),
                        "srmr": fit_i.get("srmr"),
                        "success": (it or {}).get("success"),
                        "error": (it or {}).get("error"),
                    }
                )
            invariance_df = pd.DataFrame(rows)
            invariance_estimates_df = _flatten_invariance_estimates(items)
        except Exception:
            invariance_df = pd.DataFrame([{"error": "invariance payload parse failed"}])
            invariance_estimates_df = pd.DataFrame([{"error": "invariance estimates parse failed"}])

    inv_series_models_df = None
    inv_series_lrt_df = None
    inv_series_summary_df = None
    inv_series_report_df = None
    path_summary_df = None
    integrated_summary_df = None
    model_compare_models_df = None
    model_compare_lrt_df = None
    mga_struct_df = None
    mga_struct_lrt_df = None
    if isinstance(payload.invariance_series, dict):
        # Invariance series 相关标记应写入 Meta（无论 strict/lite/降级与否），保证导出可追溯且契约稳定。
        try:
            inv_mode = payload.invariance_series.get("mode")
            inv_supported = payload.invariance_series.get("supported")
            meta_df = pd.concat(
                [
                    meta_df,
                    pd.DataFrame(
                        [
                            {"Item": "Invariance Series Mode", "Value": _safe_text(inv_mode)},
                            {"Item": "Invariance Series Supported", "Value": _safe_text(inv_supported)},
                            {"Item": "Invariance Series Message", "Value": _safe_text(payload.invariance_series.get("message"))},
                            {
                                "Item": "Invariance Series Conclusion (lite)",
                                "Value": _safe_text((payload.invariance_series_conclusion or {}).get("lite", {}).get("level"))
                                if isinstance(payload.invariance_series_conclusion, dict)
                                else "-",
                            },
                            {
                                "Item": "Invariance Series Conclusion (strict)",
                                "Value": _safe_text((payload.invariance_series_conclusion or {}).get("strict", {}).get("level"))
                                if isinstance(payload.invariance_series_conclusion, dict)
                                else "-",
                            },
                        ]
                    ),
                ],
                ignore_index=True,
            )
        except Exception:
            # Meta 只是辅助信息，不应阻断导出
            pass

    # Bootstrap（中介）协变量追溯：写入 Meta（无论是否有 items）
    if isinstance(payload.bootstrap, dict):
        try:
            covs = payload.bootstrap.get("covariates")
            cov_text = _safe_text(", ".join([str(x).strip() for x in covs]) if isinstance(covs, list) and covs else None)
            meta_df = pd.concat(
                [
                    meta_df,
                    pd.DataFrame(
                        [
                            {"Item": "Bootstrap Mediation Covariates", "Value": cov_text},
                        ]
                    ),
                ],
                ignore_index=True,
            )
        except Exception:
            pass

    if isinstance(payload.invariance_series, dict) and payload.invariance_series.get("models"):
        try:
            mrows = []
            for m in payload.invariance_series.get("models") or []:
                fitv = (m or {}).get("fit") or {}
                eq = (m or {}).get("group_equal")
                mrows.append(
                    {
                        "model": (m or {}).get("model"),
                        "group_equal": ", ".join(eq) if isinstance(eq, list) else eq,
                        "chi2": _safe_na(fitv.get("chi2")),
                        "df": _safe_na(fitv.get("df")),
                        "cfi": _safe_na(fitv.get("cfi")),
                        "tli": _safe_na(fitv.get("tli")),
                        "rmsea": _safe_na(fitv.get("rmsea")),
                        "srmr": _safe_na(fitv.get("srmr")),
                        "converged": _safe_na((m or {}).get("converged")),
                    }
                )
            inv_series_models_df = pd.DataFrame(mrows)
            # 固定列顺序（论文表格友好）
            cols = ["model", "group_equal", "chi2", "df", "cfi", "tli", "rmsea", "srmr", "converged"]
            inv_series_models_df = inv_series_models_df.reindex(columns=cols)
        except Exception:
            inv_series_models_df = pd.DataFrame([{"error": "invariance_series models parse failed"}])
    if isinstance(payload.invariance_series, dict) and payload.invariance_series.get("comparisons"):
        try:
            crows = []
            for c in payload.invariance_series.get("comparisons") or []:
                if not isinstance(c, dict):
                    crows.append(
                        {
                            "from": "NA",
                            "to": "NA",
                            "ok": "NA",
                            "chi2_diff": "NA",
                            "df_diff": "NA",
                            "p_value": "NA",
                            "delta_cfi": "NA",
                            "delta_rmsea": "NA",
                            "note": "invalid row",
                        }
                    )
                    continue
                crows.append(
                    {
                        "from": c.get("from"),
                        "to": c.get("to"),
                        "ok": c.get("ok"),
                        "chi2_diff": _safe_na(c.get("chi2_diff")),
                        "df_diff": _safe_na(c.get("df_diff")),
                        "p_value": _safe_na(c.get("p_value")),
                        "delta_cfi": _safe_na(c.get("delta_cfi")),
                        "delta_rmsea": _safe_na(c.get("delta_rmsea")),
                        "note": c.get("note"),
                    }
                )
            inv_series_lrt_df = pd.DataFrame(crows)
            cols = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "delta_cfi", "delta_rmsea", "note"]
            inv_series_lrt_df = inv_series_lrt_df.reindex(columns=cols)
        except Exception:
            inv_series_lrt_df = pd.DataFrame([{"error": "invariance_series comparisons parse failed"}])

    # UI 生成的摘要（lite/strict 结论、建议行）也写入导出（可追溯/可复制）
    try:
        inv_series_summary_df = _build_invariance_series_summary_df(
            summary_lines=payload.invariance_series_summary_lines,
            conclusion=payload.invariance_series_conclusion,
        )
    except Exception:
        inv_series_summary_df = pd.DataFrame([{"error": "invariance_series summary parse failed"}])

    # UI 生成的“报告模板（lite/strict）”写入导出（xlsx 可追溯/可复制）
    try:
        inv_series_report_df = _build_invariance_series_report_df(payload.invariance_series_report)
    except Exception:
        inv_series_report_df = pd.DataFrame([{"error": "invariance_series report parse failed"}])

    # UI 生成的“路径显著性汇总/摘要段落”写入导出（可追溯/可复制）
    try:
        path_summary_df = _build_path_summary_df(
            rows=payload.path_summary_rows,
            summary_lines=payload.path_summary_lines,
            report_text=payload.path_summary_report_text,
        )
    except Exception:
        path_summary_df = pd.DataFrame([{"error": "path_summary parse failed"}])

    try:
        ip = str(payload.integrated_paper_paragraph or "").strip()
        if ip:
            integrated_summary_df = pd.DataFrame([{"integrated_result_paragraph": ip}])
    except Exception:
        integrated_summary_df = pd.DataFrame([{"error": "integrated paragraph parse failed"}])

    if isinstance(payload.model_comparison, dict):
        try:
            mc_mode = payload.model_comparison.get("mode")
            mc_supported = payload.model_comparison.get("supported")
            meta_df = pd.concat(
                [
                    meta_df,
                    pd.DataFrame(
                        [
                            {"Item": "Model Compare Mode", "Value": _safe_text(mc_mode)},
                            {"Item": "Model Compare Supported", "Value": _safe_text(mc_supported)},
                            {"Item": "Model Compare Message", "Value": _safe_text(payload.model_comparison.get("message"))},
                        ]
                    ),
                ],
                ignore_index=True,
            )
        except Exception:
            pass

    if isinstance(payload.model_comparison, dict) and payload.model_comparison.get("models"):
        try:
            rows = []
            for m in payload.model_comparison.get("models") or []:
                fitv = (m or {}).get("fit") or {}
                rows.append(
                    {
                        "label": (m or {}).get("label"),
                        "aic": _safe_na(fitv.get("aic")),
                        "bic": _safe_na(fitv.get("bic")),
                        "chi2": _safe_na(fitv.get("chi2")),
                        "df": _safe_na(fitv.get("df")),
                        "cfi": _safe_na(fitv.get("cfi")),
                        "tli": _safe_na(fitv.get("tli")),
                        "rmsea": _safe_na(fitv.get("rmsea")),
                        "srmr": _safe_na(fitv.get("srmr")),
                    }
                )
            model_compare_models_df = pd.DataFrame(rows)
            cols = ["label", "aic", "bic", "chi2", "df", "cfi", "tli", "rmsea", "srmr"]
            model_compare_models_df = model_compare_models_df.reindex(columns=cols)
        except Exception:
            model_compare_models_df = pd.DataFrame([{"error": "model_comparison models parse failed"}])

    if isinstance(payload.model_comparison, dict) and payload.model_comparison.get("comparison"):
        try:
            c = payload.model_comparison.get("comparison") or {}
            model_compare_lrt_df = pd.DataFrame(
                [
                    {
                        "from": c.get("from"),
                        "to": c.get("to"),
                        "ok": c.get("ok"),
                        "chi2_diff": _safe_na(c.get("chi2_diff")),
                        "df_diff": _safe_na(c.get("df_diff")),
                        "p_value": _safe_na(c.get("p_value")),
                        "note": c.get("note"),
                    }
                ]
            )
            cols = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "note"]
            model_compare_lrt_df = model_compare_lrt_df.reindex(columns=cols)
        except Exception:
            model_compare_lrt_df = pd.DataFrame([{"error": "model_comparison comparison parse failed"}])

    if isinstance(payload.mga_structural, dict):
        try:
            mg_mode = payload.mga_structural.get("mode")
            mg_supported = payload.mga_structural.get("supported")
            meta_df = pd.concat(
                [
                    meta_df,
                    pd.DataFrame(
                        [
                            {"Item": "MGA Structural Mode", "Value": _safe_text(mg_mode)},
                            {"Item": "MGA Structural Supported", "Value": _safe_text(mg_supported)},
                            {"Item": "MGA Structural Message", "Value": _safe_text(payload.mga_structural.get("message"))},
                        ]
                    ),
                ],
                ignore_index=True,
            )
        except Exception:
            pass

    if isinstance(payload.mga_structural, dict):
        try:
            mga_struct_df, mga_struct_lrt_df = _build_mga_structural_dfs(payload.mga_structural)
        except Exception:
            mga_struct_df = pd.DataFrame([{"error": "mga_structural parse failed"}])
            mga_struct_lrt_df = pd.DataFrame([{"error": "mga_structural parse failed"}])

    latent_interaction_df = None
    if isinstance(payload.latent_interaction, dict) and (
        payload.latent_interaction.get("matching") or payload.latent_interaction.get("structural_for_y")
    ):
        try:
            latent_interaction_df = _flatten_latent_interaction_rows(payload.latent_interaction)
        except Exception:
            latent_interaction_df = pd.DataFrame([{"error": "latent_interaction payload parse failed"}])

    try:
        buf = BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            fit_df.to_excel(writer, sheet_name="Fit_Indices", index=False)
            if not est_df.empty:
                est_df.to_excel(writer, sheet_name="Estimates", index=False)
            syntax_df.to_excel(writer, sheet_name="Model_Syntax", index=False)
            meta_df.to_excel(writer, sheet_name="Meta", index=False)
            if path_summary_df is not None and not path_summary_df.empty:
                path_summary_df.to_excel(writer, sheet_name="Path_Summary", index=False)
            if integrated_summary_df is not None and not integrated_summary_df.empty:
                integrated_summary_df.to_excel(writer, sheet_name="Integrated_Summary", index=False)
            if bootstrap_df is not None and not bootstrap_df.empty:
                bootstrap_df.to_excel(writer, sheet_name="Bootstrap", index=False)
            if moderation_df is not None and not moderation_df.empty:
                moderation_df.to_excel(writer, sheet_name="Moderation", index=False)
            if moderated_mediation_df is not None and not moderated_mediation_df.empty:
                moderated_mediation_df.to_excel(writer, sheet_name="ModMediation", index=False)
            if invariance_df is not None and not invariance_df.empty:
                invariance_df.to_excel(writer, sheet_name="Invariance", index=False)
            if invariance_estimates_df is not None and not invariance_estimates_df.empty:
                invariance_estimates_df.to_excel(writer, sheet_name="Invariance_Estimates", index=False)
            if inv_series_models_df is not None and not inv_series_models_df.empty:
                inv_series_models_df.to_excel(writer, sheet_name="Invariance_Series_Models", index=False)
            if inv_series_lrt_df is not None and not inv_series_lrt_df.empty:
                inv_series_lrt_df.to_excel(writer, sheet_name="Invariance_Series_LRT", index=False)
            if inv_series_summary_df is not None and not inv_series_summary_df.empty:
                inv_series_summary_df.to_excel(writer, sheet_name="Invariance_Series_Summary", index=False)
            if inv_series_report_df is not None and not inv_series_report_df.empty:
                inv_series_report_df.to_excel(writer, sheet_name="Invariance_Series_Report", index=False)
            if model_compare_models_df is not None and not model_compare_models_df.empty:
                model_compare_models_df.to_excel(writer, sheet_name="Model_Compare_Models", index=False)
            if model_compare_lrt_df is not None and not model_compare_lrt_df.empty:
                model_compare_lrt_df.to_excel(writer, sheet_name="Model_Compare_LRT", index=False)
            if mga_struct_df is not None and not mga_struct_df.empty:
                mga_struct_df.to_excel(writer, sheet_name="MGA_Structural_Path", index=False)
            if mga_struct_lrt_df is not None and not mga_struct_lrt_df.empty:
                mga_struct_lrt_df.to_excel(writer, sheet_name="MGA_Structural_LRT", index=False)
            if latent_interaction_df is not None and not latent_interaction_df.empty:
                latent_interaction_df.to_excel(writer, sheet_name="Latent_Interaction", index=False)
        content = buf.getvalue()
    except Exception as e:
        _raise_export_error(500, f"APA xlsx 生成失败：{str(e)}", "EXPORT_APA_BUILD_FAILED")

    filename = _build_filename(payload, "apa_table.xlsx")
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )


@router.post("/apa-table-docx")
def export_apa_table_docx(payload: ExportPayload):
    """
    导出 APA 7th 基础报告 (.docx)
    """
    df, syntax = _validate_export_payload(payload)
    template_key, template = _resolve_export_template(payload)
    fit = payload.fit_indices or {}
    estimates = payload.estimates or []
    bootstrap = payload.bootstrap or None
    moderation = payload.moderation or None
    moderated_mediation = payload.moderated_mediation or None
    invariance = payload.invariance or None
    invariance_series = payload.invariance_series or None
    model_comparison = payload.model_comparison or None
    mga_structural = payload.mga_structural or None
    latent_interaction = payload.latent_interaction or None

    try:
        doc = Document()
        doc.add_heading("OpenSEM APA Report (Phase 1)", level=1)
        doc.add_paragraph(f"N Rows (raw): {int(len(df))}")
        doc.add_paragraph(f"N Cols: {int(len(df.columns))}")
        doc.add_paragraph(f"Export template key: {template_key}")
        doc.add_paragraph(f"Export template name: {_safe_text(template.get('name'))}")
        doc.add_paragraph(f"Estimator: {_safe_text(payload.estimator)}")
        doc.add_paragraph(f"Missing strategy: {_safe_text(payload.missing_strategy)}")
        doc.add_paragraph(f"Fit ran at (client): {_safe_text(payload.fit_ran_at)}")
        if payload.fit_stale:
            doc.add_paragraph("WARNING: Export is STALE — syntax changed after last fit; interpret with caution.")
        doc.add_paragraph(f"Syntax tag (client): {_safe_text(payload.syntax_tag)}")
        doc.add_paragraph(f"Last fit syntax tag (client): {_safe_text(payload.last_fit_syntax_tag)}")
        doc.add_paragraph(f"MI adoption count: {_safe_text(payload.mi_adoption_count)}")
        doc.add_paragraph(f"MI ignored count: {_safe_text(payload.mi_ignored_count)}")
        doc.add_paragraph(f"MI risk filter: {_safe_text(payload.mi_risk_filter)}")
        doc.add_paragraph(f"MI adoption bucket tag: {_safe_text(payload.mi_adoption_bucket_tag)}")
        doc.add_paragraph(f"MI adoption first at: {_safe_text(payload.mi_adoption_first_at)}")
        doc.add_paragraph(f"MI adoption last at: {_safe_text(payload.mi_adoption_last_at)}")
        doc.add_paragraph(f"MI adoption last id: {_safe_text(payload.mi_adoption_last_id)}")
        doc.add_paragraph(f"Syntax fingerprint (sha256/12): {_syntax_fingerprint(syntax)}")

        doc.add_heading("Fit Indices", level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Indicator"
        hdr[1].text = "Value"
        for k in ["chi2", "chi2_df", "rmsea", "srmr", "cfi", "tli", "status"]:
            row = table.add_row().cells
            row[0].text = k
            row[1].text = _safe_text(fit.get(k))

        if estimates:
            doc.add_heading("Parameter Estimates", level=2)
            cols = ["LHS", "Op", "RHS", "Estimate", "Beta(Std)", "S.E.", "C.R.", "p"]
            t2 = doc.add_table(rows=1, cols=len(cols))
            hdr2 = t2.rows[0].cells
            for i, c in enumerate(cols):
                hdr2[i].text = c
            estimate_rows, estimate_omitted = _truncate_docx_rows(estimates)
            for r in estimate_rows:
                row = t2.add_row().cells
                row[0].text = _safe_text(r.get("lval"))
                row[1].text = _safe_text(r.get("op"))
                row[2].text = _safe_text(r.get("rval"))
                row[3].text = _safe_text(r.get("estimate"))
                row[4].text = _safe_text(r.get("est_std"))
                row[5].text = _safe_text(r.get("std_err"))
                row[6].text = _safe_text(r.get("z_value"))
                row[7].text = _safe_text(r.get("p_value"))
            if estimate_omitted > 0:
                doc.add_paragraph(f"（Parameter Estimates 已截断展示前 {len(estimate_rows)} 行；其余 {estimate_omitted} 行请见 xlsx 的 Estimates）")

        # 1.3：综合结果摘要（UI 拼装）
        try:
            ip = str(payload.integrated_paper_paragraph or "").strip()
            if ip:
                doc.add_heading("Integrated results summary (UI)", level=2)
                for line in ip.splitlines():
                    if line.strip() == "":
                        doc.add_paragraph("")
                    else:
                        doc.add_paragraph(line)
        except Exception:
            pass

        # Wave 2 / P1：路径显著性汇总（UI 侧生成）——用于论文写作直接复制/追溯
        try:
            lines = payload.path_summary_lines or []
            report_text = str(payload.path_summary_report_text or "").strip()
            rows = payload.path_summary_rows or []
            has_lines = isinstance(lines, list) and any(str(x).strip() for x in lines)
            has_report = bool(report_text)
            has_rows = isinstance(rows, list) and any(isinstance(x, dict) for x in rows)
            if has_lines or has_report or has_rows:
                doc.add_heading("Path significance summary (UI)", level=2)
                if has_lines:
                    for x in lines:
                        s = str(x or "").strip()
                        if s:
                            doc.add_paragraph(s)
                if has_report:
                    doc.add_heading("Path report text (UI)", level=3)
                    for line in report_text.splitlines():
                        if line.strip() == "":
                            doc.add_paragraph("")
                        else:
                            doc.add_paragraph(line)
                if has_rows:
                    doc.add_heading("Path summary table (UI)", level=3)
                    cols_p = ["predictor", "outcome", "op", "beta", "p_value", "significant", "direction", "note"]
                    tp = doc.add_table(rows=1, cols=len(cols_p))
                    hp = tp.rows[0].cells
                    for i, c in enumerate(cols_p):
                        hp[i].text = c
                    for r in rows:
                        if not isinstance(r, dict):
                            continue
                        row = tp.add_row().cells
                        row[0].text = _safe_text(r.get("predictor"))
                        row[1].text = _safe_text(r.get("outcome"))
                        row[2].text = _safe_text(r.get("op") or "~")
                        row[3].text = _safe_text(r.get("beta"))
                        row[4].text = _safe_text(r.get("p_value"))
                        row[5].text = _safe_text(r.get("significant"))
                        row[6].text = _safe_text(r.get("direction"))
                        row[7].text = _safe_text(r.get("note"))
        except Exception:
            pass

        doc.add_heading("Model Syntax", level=2)
        for line in syntax.splitlines():
            clean = line.strip()
            if clean:
                doc.add_paragraph(clean)

        if isinstance(bootstrap, dict) and bootstrap.get("items"):
            doc.add_heading("Bootstrap (Mediation)", level=2)
            doc.add_paragraph(_safe_text(bootstrap.get("task")) if bootstrap.get("task") else "bootstrap_mediation")
            try:
                covs = bootstrap.get("covariates")
                cov_text = ", ".join([str(x).strip() for x in covs]) if isinstance(covs, list) and covs else ""
                if cov_text.strip():
                    doc.add_paragraph(f"Covariates (main effects): {cov_text}")
            except Exception:
                pass
            items = bootstrap.get("items") or []
            items, bootstrap_omitted = _truncate_docx_rows(items)
            cols = ["effect_type", "label", "path", "indirect_point", "ci_level", "percentile_lo", "percentile_hi", "bc_lo", "bc_hi", "n_boot_valid", "note"]
            t3 = doc.add_table(rows=1, cols=len(cols))
            hdr3 = t3.rows[0].cells
            for i, c in enumerate(cols):
                hdr3[i].text = c
            for r in items:
                ci = (r or {}).get("ci") or {}
                perc = (ci or {}).get("percentile") or {}
                bc = (ci or {}).get("bc") or {}
                sequence = (r or {}).get("sequence")
                row = t3.add_row().cells
                row[0].text = _safe_text((r or {}).get("effect_type"))
                row[1].text = _safe_text((r or {}).get("label"))
                row[2].text = _safe_text((r or {}).get("path_label") or (" → ".join(sequence) if isinstance(sequence, list) else None))
                row[3].text = _safe_text((r or {}).get("indirect_point"))
                row[4].text = _safe_text((ci or {}).get("ci_level"))
                row[5].text = _safe_text(perc.get("lo"))
                row[6].text = _safe_text(perc.get("hi"))
                row[7].text = _safe_text(bc.get("lo"))
                row[8].text = _safe_text(bc.get("hi"))
                row[9].text = _safe_text((ci or {}).get("n_boot_valid"))
                row[10].text = _safe_text((r or {}).get("note"))
            if bootstrap_omitted > 0:
                doc.add_paragraph(f"（Bootstrap 表已截断展示前 {len(items)} 行；其余 {bootstrap_omitted} 行请见 xlsx 的 Bootstrap）")

        if isinstance(latent_interaction, dict) and (
            latent_interaction.get("matching") or latent_interaction.get("structural_for_y")
        ):
            doc.add_heading("Latent interaction (MVP probe)", level=2)
            b = latent_interaction.get("boundary") or {}
            doc.add_paragraph(_safe_text(b.get("semopy")))
            doc.add_paragraph(_safe_text(b.get("lavaan")))
            doc.add_paragraph(
                f"y={_safe_text(latent_interaction.get('y'))}, f1={_safe_text(latent_interaction.get('f1'))}, f2={_safe_text(latent_interaction.get('f2'))}, N={_safe_text(latent_interaction.get('n_used'))}"
            )
            for title, key in (("Matching (F1×F2)", "matching"), ("Structural for outcome (all rhs terms)", "structural_for_y")):
                rows = latent_interaction.get(key) or []
                if not rows:
                    continue
                doc.add_heading(title, level=3)
                cols = ["lval", "op", "rval", "estimate", "est_std", "p_value"]
                tli = doc.add_table(rows=1, cols=len(cols))
                tli.style = "Table Grid"
                for i, c in enumerate(cols):
                    tli.rows[0].cells[i].text = c
                for raw in rows:
                    rr = tli.add_row().cells
                    rr[0].text = _safe_text((raw or {}).get("lval"))
                    rr[1].text = _safe_text((raw or {}).get("op"))
                    rr[2].text = _safe_text((raw or {}).get("rval"))
                    rr[3].text = _safe_text((raw or {}).get("estimate"))
                    rr[4].text = _safe_text((raw or {}).get("est_std"))
                    rr[5].text = _safe_text((raw or {}).get("p_value"))

        if isinstance(moderation, dict) and (moderation.get("coefficients") or moderation.get("simple_slopes")):
            doc.add_heading("Moderation Analysis", level=2)
            moderator = moderation.get("moderator") or {}
            doc.add_paragraph(f"Y: {_safe_text(moderation.get('outcome'))}")
            doc.add_paragraph(f"X: {_safe_text(moderation.get('predictor'))}")
            doc.add_paragraph(
                f"W: {_safe_text(moderator.get('name'))} ({_safe_text(moderator.get('type'))}), reference={_safe_text(moderator.get('reference_group'))}"
            )

            coeff_cols = ["term", "estimate", "std_err", "z_value", "p_value"]
            t_mod = doc.add_table(rows=1, cols=len(coeff_cols))
            t_mod.style = "Table Grid"
            hdr_mod = t_mod.rows[0].cells
            for i, c in enumerate(coeff_cols):
                hdr_mod[i].text = c
            coeff_rows, coeff_omitted = _truncate_docx_rows(moderation.get("coefficients") or [])
            for row_data in coeff_rows:
                row = t_mod.add_row().cells
                row[0].text = _safe_text((row_data or {}).get("term"))
                row[1].text = _safe_text((row_data or {}).get("estimate"))
                row[2].text = _safe_text((row_data or {}).get("std_err"))
                row[3].text = _safe_text((row_data or {}).get("z_value"))
                row[4].text = _safe_text((row_data or {}).get("p_value"))
            if coeff_omitted > 0:
                doc.add_paragraph(f"（Moderation coefficients 已截断展示前 {len(coeff_rows)} 行；其余 {coeff_omitted} 行请见 xlsx 的 Moderation）")

            if moderation.get("simple_slopes"):
                doc.add_heading("Simple Slopes", level=3)
                slope_cols = ["group", "slope_x", "w_value", "note"]
                t_slope = doc.add_table(rows=1, cols=len(slope_cols))
                t_slope.style = "Table Grid"
                hdr_slope = t_slope.rows[0].cells
                for i, c in enumerate(slope_cols):
                    hdr_slope[i].text = c
                for row_data in moderation.get("simple_slopes") or []:
                    row = t_slope.add_row().cells
                    row[0].text = _safe_text((row_data or {}).get("group"))
                    row[1].text = _safe_text((row_data or {}).get("slope_x"))
                    row[2].text = _safe_text((row_data or {}).get("w_value"))
                    row[3].text = _safe_text((row_data or {}).get("note"))

        if isinstance(moderated_mediation, dict) and (
            moderated_mediation.get("conditional_indirect") or moderated_mediation.get("index_moderated_mediation")
        ):
            mm_model = str(moderated_mediation.get("model") or "")
            if "14" in mm_model:
                mm_heading = "Moderated Mediation (Hayes Model 14)"
            elif "categorical" in mm_model:
                mm_heading = "Moderated Mediation (Hayes Model 7, categorical W)"
            else:
                mm_heading = "Moderated Mediation (Hayes Model 7)"
            doc.add_heading(mm_heading, level=2)
            doc.add_paragraph(_safe_text(moderated_mediation.get("note")))
            doc.add_paragraph(
                f"X={_safe_text(moderated_mediation.get('x'))}, M={_safe_text(moderated_mediation.get('m'))}, "
                f"Y={_safe_text(moderated_mediation.get('y'))}, W={_safe_text(moderated_mediation.get('w'))}"
            )
            try:
                covs = moderated_mediation.get("covariates")
                cov_text = ", ".join([str(x).strip() for x in covs]) if isinstance(covs, list) and covs else ""
                if cov_text.strip():
                    doc.add_paragraph(f"Covariates: {cov_text}")
            except Exception:
                pass
            cols_mm = ["label", "w_value", "indirect_point", "pct_lo", "pct_hi", "bc_lo", "bc_hi", "n_boot_valid"]
            tmm = doc.add_table(rows=1, cols=len(cols_mm))
            tmm.style = "Table Grid"
            hmm = tmm.rows[0].cells
            for i, c in enumerate(cols_mm):
                hmm[i].text = c
            for row_data in moderated_mediation.get("conditional_indirect") or []:
                ci = (row_data or {}).get("ci") or {}
                perc = (ci or {}).get("percentile") or {}
                bc = (ci or {}).get("bc") or {}
                row = tmm.add_row().cells
                row[0].text = _safe_text((row_data or {}).get("label"))
                row[1].text = _safe_text((row_data or {}).get("w_value"))
                row[2].text = _safe_text((row_data or {}).get("indirect_point"))
                row[3].text = _safe_text(perc.get("lo"))
                row[4].text = _safe_text(perc.get("hi"))
                row[5].text = _safe_text(bc.get("lo"))
                row[6].text = _safe_text(bc.get("hi"))
                row[7].text = _safe_text((ci or {}).get("n_boot_valid"))
            idx = moderated_mediation.get("index_moderated_mediation") or {}
            if isinstance(idx, dict) and idx:
                doc.add_heading("Index of moderated mediation", level=3)
                ci = idx.get("ci") or {}
                perc = (ci or {}).get("percentile") or {}
                bc = (ci or {}).get("bc") or {}
                doc.add_paragraph(_safe_text(idx.get("contrast")))
                doc.add_paragraph(
                    f"Point={_safe_text(idx.get('point'))}, "
                    f"Percentile CI [{_safe_text(perc.get('lo'))}, {_safe_text(perc.get('hi'))}], "
                    f"BC CI [{_safe_text(bc.get('lo'))}, {_safe_text(bc.get('hi'))}], "
                    f"n_boot_valid={_safe_text((ci or {}).get('n_boot_valid'))}"
                )

        if isinstance(model_comparison, dict):
            doc.add_heading("Model Comparison (AIC/BIC + LRT)", level=2)
            if model_comparison.get("mode") == "degraded" or model_comparison.get("supported") is False:
                doc.add_paragraph("DEGRADED: model comparison not fully supported in current environment.")
                if model_comparison.get("message"):
                    doc.add_paragraph(_safe_text(model_comparison.get("message")))
            if model_comparison.get("note"):
                doc.add_paragraph(_safe_text(model_comparison.get("note")))

            models = model_comparison.get("models") or []
            if isinstance(models, list) and models:
                cols_mc = ["label", "aic", "bic", "chi2", "df", "cfi", "tli", "rmsea", "srmr"]
                tmc = doc.add_table(rows=1, cols=len(cols_mc))
                hmc = tmc.rows[0].cells
                for i, c in enumerate(cols_mc):
                    hmc[i].text = c
                for m in models:
                    fitv = (m or {}).get("fit") or {}
                    row = tmc.add_row().cells
                    row[0].text = _safe_text((m or {}).get("label"))
                    row[1].text = _safe_text(fitv.get("aic"))
                    row[2].text = _safe_text(fitv.get("bic"))
                    row[3].text = _safe_text(fitv.get("chi2"))
                    row[4].text = _safe_text(fitv.get("df"))
                    row[5].text = _safe_text(fitv.get("cfi"))
                    row[6].text = _safe_text(fitv.get("tli"))
                    row[7].text = _safe_text(fitv.get("rmsea"))
                    row[8].text = _safe_text(fitv.get("srmr"))

            comp = model_comparison.get("comparison") or {}
            if isinstance(comp, dict) and comp:
                doc.add_heading("Nested LRT (lavTestLRT)", level=3)
                cols_lrt = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "note"]
                tlrt = doc.add_table(rows=1, cols=len(cols_lrt))
                hlrt = tlrt.rows[0].cells
                for i, c in enumerate(cols_lrt):
                    hlrt[i].text = c
                row = tlrt.add_row().cells
                row[0].text = _safe_text(comp.get("from"))
                row[1].text = _safe_text(comp.get("to"))
                row[2].text = _safe_text(comp.get("ok"))
                row[3].text = _safe_text(comp.get("chi2_diff"))
                row[4].text = _safe_text(comp.get("df_diff"))
                row[5].text = _safe_text(comp.get("p_value"))
                row[6].text = _safe_text(comp.get("note"))

        if isinstance(mga_structural, dict):
            doc.add_heading("MGA Structural Path Compare", level=2)
            if mga_structural.get("group_var"):
                doc.add_paragraph(f"Group var: {_safe_text(mga_structural.get('group_var'))}")
            path = mga_structural.get("path") or {}
            if isinstance(path, dict) and (path.get("predictor") or path.get("outcome")):
                doc.add_paragraph(f"Path: {_safe_text(path.get('predictor'))} → {_safe_text(path.get('outcome'))}")
            if mga_structural.get("mode") == "degraded" or mga_structural.get("supported") is False:
                doc.add_paragraph("DEGRADED: MGA structural compare not fully supported in current environment.")
                if mga_structural.get("message"):
                    doc.add_paragraph(_safe_text(mga_structural.get("message")))
            if mga_structural.get("note"):
                doc.add_paragraph(_safe_text(mga_structural.get("note")))

            # Wave 3：层级×多路径（items）
            if isinstance(mga_structural.get("items"), list) and mga_structural.get("items"):
                for it in mga_structural.get("items") or []:
                    if not isinstance(it, dict):
                        continue
                    level = _safe_text(it.get("level"))
                    p = it.get("path") if isinstance(it.get("path"), dict) else {}
                    predictor = _safe_text(p.get("predictor"))
                    outcome = _safe_text(p.get("outcome"))
                    if level or predictor or outcome:
                        doc.add_heading(f"{level}: {predictor} → {outcome}", level=3)

                    ge_items = it.get("group_estimates") or []
                    if isinstance(ge_items, list) and ge_items:
                        cols = ["group", "estimate", "std_all", "p_value", "note"]
                        tmg = doc.add_table(rows=1, cols=len(cols))
                        hmg = tmg.rows[0].cells
                        for i, c in enumerate(cols):
                            hmg[i].text = c
                        for ge in ge_items:
                            if not isinstance(ge, dict):
                                continue
                            coef = ge.get("coef") if isinstance(ge.get("coef"), dict) else {}
                            row = tmg.add_row().cells
                            row[0].text = _safe_text(ge.get("group"))
                            row[1].text = _safe_text(ge.get("estimate") if ge.get("estimate") is not None else coef.get("estimate"))
                            row[2].text = _safe_text(ge.get("std_all") if ge.get("std_all") is not None else coef.get("std_all"))
                            row[3].text = _safe_text(ge.get("p_value") if ge.get("p_value") is not None else coef.get("p_value"))
                            row[4].text = _safe_text(ge.get("error") if ge.get("success") is False else (ge.get("note") or ""))

                    comp = it.get("comparison") or {}
                    if isinstance(comp, dict) and comp:
                        cols = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "note"]
                        t2 = doc.add_table(rows=1, cols=len(cols))
                        h = t2.rows[0].cells
                        for i, c in enumerate(cols):
                            h[i].text = c
                        r = t2.add_row().cells
                        r[0].text = _safe_text(comp.get("from"))
                        r[1].text = _safe_text(comp.get("to"))
                        r[2].text = _safe_text(comp.get("ok"))
                        r[3].text = _safe_text(comp.get("chi2_diff"))
                        r[4].text = _safe_text(comp.get("df_diff"))
                        r[5].text = _safe_text(comp.get("p_value"))
                        r[6].text = _safe_text(comp.get("note"))

            else:
                # 旧结构：单路径
                items = mga_structural.get("group_estimates") or []
                if isinstance(items, list) and items:
                    cols = ["group", "estimate", "std_all", "p_value", "note"]
                    tmg = doc.add_table(rows=1, cols=len(cols))
                    hmg = tmg.rows[0].cells
                    for i, c in enumerate(cols):
                        hmg[i].text = c
                    for it in items:
                        if not isinstance(it, dict):
                            continue
                        coef = it.get("coef") if isinstance(it.get("coef"), dict) else {}
                        row = tmg.add_row().cells
                        row[0].text = _safe_text(it.get("group"))
                        row[1].text = _safe_text(it.get("estimate") if it.get("estimate") is not None else coef.get("estimate"))
                        row[2].text = _safe_text(it.get("std_all") if it.get("std_all") is not None else coef.get("std_all"))
                        row[3].text = _safe_text(it.get("p_value") if it.get("p_value") is not None else coef.get("p_value"))
                        row[4].text = _safe_text(it.get("error") if it.get("success") is False else (it.get("note") or ""))

                comp = mga_structural.get("comparison") or {}
                if isinstance(comp, dict) and comp:
                    doc.add_heading("MGA Nested LRT", level=3)
                    cols = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "note"]
                    t2 = doc.add_table(rows=1, cols=len(cols))
                    h = t2.rows[0].cells
                    for i, c in enumerate(cols):
                        h[i].text = c
                    r = t2.add_row().cells
                    r[0].text = _safe_text(comp.get("from"))
                    r[1].text = _safe_text(comp.get("to"))
                    r[2].text = _safe_text(comp.get("ok"))
                    r[3].text = _safe_text(comp.get("chi2_diff"))
                    r[4].text = _safe_text(comp.get("df_diff"))
                    r[5].text = _safe_text(comp.get("p_value"))
                    r[6].text = _safe_text(comp.get("note"))

        if isinstance(invariance, dict) and invariance.get("items"):
            doc.add_heading("Multi-group (Configural)", level=2)
            doc.add_paragraph(_safe_text(invariance.get("group_var")))
            items = invariance.get("items") or []
            cols = ["group", "n_used", "status", "chi2_df", "rmsea", "cfi", "tli", "srmr", "success", "error"]
            t4 = doc.add_table(rows=1, cols=len(cols))
            hdr4 = t4.rows[0].cells
            for i, c in enumerate(cols):
                hdr4[i].text = c
            for r in items:
                fit_i = (r or {}).get("fit_indices") or {}
                row = t4.add_row().cells
                row[0].text = _safe_text((r or {}).get("group"))
                row[1].text = _safe_text((r or {}).get("n_used"))
                row[2].text = _safe_text(fit_i.get("status"))
                row[3].text = _safe_text(fit_i.get("chi2_df"))
                row[4].text = _safe_text(fit_i.get("rmsea"))
                row[5].text = _safe_text(fit_i.get("cfi"))
                row[6].text = _safe_text(fit_i.get("tli"))
                row[7].text = _safe_text(fit_i.get("srmr"))
                row[8].text = _safe_text((r or {}).get("success"))
                row[9].text = _safe_text((r or {}).get("error"))

            # 导出增强：各组参数表（截断版，避免 docx 过大；完整表在 xlsx 的 Invariance_Estimates）
            try:
                shown_groups = 0
                for it in items:
                    if not isinstance(it, dict):
                        continue
                    est = it.get("estimates") or []
                    if not isinstance(est, list) or not est:
                        continue
                    shown_groups += 1
                    if shown_groups > 3:
                        doc.add_paragraph("（其余组别的参数表已在 xlsx 的 Invariance_Estimates 中提供）")
                        break
                    doc.add_heading(f"Group estimates (truncated): {_safe_text(it.get('group'))}", level=3)
                    cols_e = ["lval", "op", "rval", "estimate", "est_std", "std_err", "z_value", "p_value"]
                    te = doc.add_table(rows=1, cols=len(cols_e))
                    he = te.rows[0].cells
                    for i, c in enumerate(cols_e):
                        he[i].text = c
                    for r in est[:15]:
                        if not isinstance(r, dict):
                            continue
                        row = te.add_row().cells
                        row[0].text = _safe_text(r.get("lval"))
                        row[1].text = _safe_text(r.get("op"))
                        row[2].text = _safe_text(r.get("rval"))
                        row[3].text = _safe_text(r.get("estimate"))
                        row[4].text = _safe_text(r.get("est_std"))
                        row[5].text = _safe_text(r.get("std_err"))
                        row[6].text = _safe_text(r.get("z_value"))
                        row[7].text = _safe_text(r.get("p_value"))
            except Exception:
                pass

        if isinstance(invariance_series, dict):
            doc.add_heading("Multi-group invariance sequence (lavaan)", level=2)
            if invariance_series.get("group_var"):
                doc.add_paragraph(_safe_text(invariance_series.get("group_var")))

            if invariance_series.get("mode") == "degraded" or invariance_series.get("supported") is False:
                doc.add_paragraph("DEGRADED: invariance series not supported in current environment.")
                doc.add_paragraph(
                    _safe_text(invariance_series.get("message"))
                    or "Invariance series degraded / not supported in current environment."
                )
            if invariance_series.get("note"):
                doc.add_paragraph(_safe_text(invariance_series.get("note")))

            # UI 侧摘要/报告段落（lite/strict）——用于论文写作直接复制
            try:
                lines = payload.invariance_series_summary_lines or []
                if isinstance(lines, list) and any(str(x).strip() for x in lines):
                    doc.add_heading("Invariance summary (UI)", level=3)
                    for x in lines:
                        s = str(x or "").strip()
                        if s:
                            doc.add_paragraph(s)
            except Exception:
                pass

            try:
                report = payload.invariance_series_report or {}
                if isinstance(report, dict) and (str(report.get("lite") or "").strip() or str(report.get("strict") or "").strip()):
                    doc.add_heading("Invariance report text (UI)", level=3)
                    lite_text = str(report.get("lite") or "").strip()
                    strict_text = str(report.get("strict") or "").strip()
                    if lite_text:
                        doc.add_paragraph("Lite (up to scalar):")
                        for line in lite_text.splitlines():
                            if line.strip() == "":
                                doc.add_paragraph("")
                            else:
                                doc.add_paragraph(line)
                    if strict_text:
                        doc.add_paragraph("Strict (up to strict):")
                        for line in strict_text.splitlines():
                            if line.strip() == "":
                                doc.add_paragraph("")
                            else:
                                doc.add_paragraph(line)
            except Exception:
                pass

            items_m = invariance_series.get("models") or []
            if items_m:
                # 与 xlsx 契约一致：包含 converged 列
                cols_m = ["model", "group_equal", "chi2", "df", "cfi", "tli", "rmsea", "srmr", "converged"]
                tm = doc.add_table(rows=1, cols=len(cols_m))
                hm = tm.rows[0].cells
                for i, c in enumerate(cols_m):
                    hm[i].text = c
                for r in items_m:
                    fitv = (r or {}).get("fit") or {}
                    eq = (r or {}).get("group_equal")
                    row = tm.add_row().cells
                    row[0].text = _safe_text((r or {}).get("model"))
                    row[1].text = _safe_text(", ".join(eq) if isinstance(eq, list) else eq)
                    row[2].text = _safe_text(fitv.get("chi2"))
                    row[3].text = _safe_text(fitv.get("df"))
                    row[4].text = _safe_text(fitv.get("cfi"))
                    row[5].text = _safe_text(fitv.get("tli"))
                    row[6].text = _safe_text(fitv.get("rmsea"))
                    row[7].text = _safe_text(fitv.get("srmr"))
                    row[8].text = _safe_text((r or {}).get("converged"))

            items_c = invariance_series.get("comparisons") or []
            if items_c:
                doc.add_heading("Invariance LRT / ΔCFI", level=3)
                # 与 xlsx 契约一致：包含 ok/note 列
                cols_c = ["from", "to", "ok", "chi2_diff", "df_diff", "p_value", "delta_cfi", "delta_rmsea", "note"]
                tc = doc.add_table(rows=1, cols=len(cols_c))
                hc = tc.rows[0].cells
                for i, c in enumerate(cols_c):
                    hc[i].text = c
                for c in items_c:
                    if not isinstance(c, dict):
                        continue
                    row = tc.add_row().cells
                    row[0].text = _safe_text(c.get("from"))
                    row[1].text = _safe_text(c.get("to"))
                    row[2].text = _safe_text(c.get("ok"))
                    row[3].text = _safe_text(c.get("chi2_diff"))
                    row[4].text = _safe_text(c.get("df_diff"))
                    row[5].text = _safe_text(c.get("p_value"))
                    row[6].text = _safe_text(c.get("delta_cfi"))
                    row[7].text = _safe_text(c.get("delta_rmsea"))
                    row[8].text = _safe_text(c.get("note"))

        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()
    except Exception as e:
        _raise_export_error(500, f"APA docx 生成失败：{str(e)}", "EXPORT_APA_DOCX_BUILD_FAILED")

    filename = _build_filename(payload, "apa_table.docx")
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.post("/lavaan-syntax")
def export_lavaan_syntax(payload: ExportPayload):
    """
    导出 lavaan 语法
    """
    _df, syntax = _validate_export_payload(payload)

    filename = _build_filename(payload, "model.lavaan.txt")
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(content=syntax.encode("utf-8"), media_type="text/plain; charset=utf-8", headers=headers)
